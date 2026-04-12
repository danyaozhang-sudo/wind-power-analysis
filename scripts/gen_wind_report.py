#!/usr/bin/env python3
"""
风电项目 Word 报告生成脚本（v2固化标准）。

输入: model_data.json（含完整财务数据）
输出: Word 报告文件

用法:
    python gen_wind_report.py <json_path> <charts_dir> <logo_path> <qr_path> <output_docx>
"""
import sys, json, os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLUE  = RGBColor(44, 82, 130)
WHITE = RGBColor(255, 255, 255)
DARK  = RGBColor(68, 68, 68)
GRAY  = RGBColor(170, 170, 170)

def set_cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, size=8, color=DARK, align='center', font='微软雅黑'):
    cell.text = ''
    pg = cell.paragraphs[0]
    pg.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == 'center' else (
        WD_ALIGN_PARAGRAPH.LEFT if align == 'left' else WD_ALIGN_PARAGRAPH.RIGHT)
    r = pg.add_run(text)
    r.font.bold = bold; r.font.size = Pt(size); r.font.color.rgb = color; r.font.name = font
    if r._element.rPr is None: r._element.rPr = r._element.get_or_add_rPr()
    r._element.rPr.rFonts.set(qn('w:eastAsia'), font)

def h1(doc, num, title):
    pg = doc.add_paragraph()
    pg.paragraph_format.space_before = Pt(14); pg.paragraph_format.space_after = Pt(4)
    r = pg.add_run(f'第{num}章  {title}')
    r.font.bold = True; r.font.size = Pt(13); r.font.color.rgb = BLUE; r.font.name = '微软雅黑'
    if r._element.rPr is None: r._element.rPr = r._element.get_or_add_rPr()
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def h2(doc, title):
    pg = doc.add_paragraph()
    pg.paragraph_format.space_before = Pt(8); pg.paragraph_format.space_after = Pt(2)
    r = pg.add_run(title)
    r.font.bold = True; r.font.size = Pt(10); r.font.color.rgb = BLUE; r.font.name = '微软雅黑'
    if r._element.rPr is None: r._element.rPr = r._element.get_or_add_rPr()
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def note(doc, txt):
    pg = doc.add_paragraph()
    pg.paragraph_format.space_before = Pt(2); pg.paragraph_format.space_after = Pt(2)
    r = pg.add_run('公式：' + txt)
    r.font.size = Pt(7.5); r.font.color.rgb = GRAY; r.font.name = '微软雅黑'
    if r._element.rPr is None: r._element.rPr = r._element.get_or_add_rPr()
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def make_table(doc, hdrs, rows, widths, sz=8):
    n_cols = len(hdrs)
    tbl = doc.add_table(rows=1+len(rows), cols=n_cols)
    tbl.style = 'Table Grid'; tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(hdrs):
        cell = tbl.rows[0].cells[i]; set_cell_bg(cell, '2B528A'); set_cell_text(cell, h, bold=True, size=sz-1, color=WHITE)
    for ri, row in enumerate(rows):
        fill = 'F5F5F5' if ri%2==0 else 'FFFFFF'
        for ci, (cell, val) in enumerate(zip(tbl.rows[ri+1].cells, row)):
            set_cell_bg(cell, fill); set_cell_text(cell, str(val), size=sz)
    for i, w in enumerate(widths):
        for row in tbl.rows: row.cells[i].width = Cm(w)
    doc.add_paragraph()

def dscr_status(dscr):
    if dscr is None: return '已还清'
    elif dscr > 1.2: return '健康'
    elif dscr > 1.0: return '预警'
    else: return '危险'

def generate_report(json_path, charts_dir, logo_path, qr_path, output_path):
    D = json.load(open(json_path, encoding='utf-8'))
    m = D['metrics']; p = D['params']
    base_irr = m['irr_equity']
    os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)

    doc = Document()

    # ========== 封面 ==========
    doc.add_paragraph()
    pg_logo = doc.add_paragraph()
    pg_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_logo = pg_logo.add_run(); r_logo.add_picture(logo_path, width=Inches(2.0))

    doc.add_paragraph()
    pg_title = doc.add_paragraph(); pg_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = pg_title.add_run(D['project_name'])
    r.font.size = Pt(24); r.font.bold = True; r.font.color.rgb = BLUE; r.font.name = '微软雅黑'
    if r._element.rPr is None: r._element.rPr = r._element.get_or_add_rPr()
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    pg_sub = doc.add_paragraph(); pg_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = pg_sub.add_run('财务分析报告')
    r.font.size = Pt(16); r.font.color.rgb = BLUE; r.font.name = '微软雅黑'
    if r._element.rPr is None: r._element.rPr = r._element.get_or_add_rPr()
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_paragraph()
    kpis = [
        ('全投资 IRR', '%.2f%%' % m['irr_full']),
        ('资本金 IRR', '%.2f%%' % m['irr_equity']),
        ('全投资 NPV', '%.0f 万元' % m['npv_full']),
        ('资本金 NPV', '%.0f 万元' % m['npv_equity']),
        ('ROE（年均）', '%.1f%%' % m['roe']),
        ('ROI（累计）', '%.1f%%' % m['roi']),
    ]
    tbl_kpi = doc.add_table(rows=2, cols=6); tbl_kpi.style = 'Table Grid'; tbl_kpi.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (lbl, val) in enumerate(kpis):
        c0 = tbl_kpi.rows[0].cells[i]; set_cell_bg(c0, '2B528A'); set_cell_text(c0, lbl, bold=True, size=8, color=WHITE)
        c1 = tbl_kpi.rows[1].cells[i]; set_cell_bg(c1, 'F5F5F5'); set_cell_text(c1, val, bold=True, size=12, color=BLUE)
    for i in range(6):
        for row in tbl_kpi.rows: row.cells[i].width = Cm(3.0)
    doc.add_paragraph()

    h2(doc, '项目基本信息')
    basic_rows = [
        ('装机容量', '%d MW' % p['capacity']),
        ('EPC单价', '%.1f 元/W' % p['epc_price']),
        ('年利用小时', '%d h（限电%+.0f%%后）' % (p['util_hours'], p['curtailment']*100)),
        ('CFD电价', '%.2f 元/kWh（前%d年）' % (p['cfd_price'], p['cfd_years'])),
        ('市场化电价', '%.2f 元/kWh（%d年后）' % (p['post_cfd_price'], p['cfd_years'])),
        ('总投资', '%.0f 万元（%.2f亿元）' % (p['total_inv'], p['total_inv']/10000)),
        ('资本金（20%%）', '%.0f 万元' % p['capital']),
        ('贷款（80%%）', '%.0f 万元' % p['loan']),
        ('贷款利率', '%.1f%%（含税）' % (p['loan_rate']*100)),
        ('融资期限', '%d 年（等额本金）' % p['loan_years']),
        ('年发电量', '%.0f 万kWh/年' % p['annual_power']),
        ('残值', '%.0f 万元（总投资x3%%）' % p['residual_value']),
        ('年折旧', '%.0f 万元/年' % p['annual_depreciation']),
    ]
    make_table(doc, ['参数', '数值'], basic_rows, [5.0, 9.0], sz=9)
    doc.add_paragraph()

    pg_org = doc.add_paragraph(); pg_org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = pg_org.add_run('JanoPower经济研究院 · 投资部    2026年4月')
    r.font.size = Pt(10); r.font.color.rgb = GRAY; r.font.name = '微软雅黑'
    if r._element.rPr is None: r._element.rPr = r._element.get_or_add_rPr()
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    pg_qr = doc.add_paragraph(); pg_qr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_qr = pg_qr.add_run(); r_qr.add_picture(qr_path, width=Inches(1.5))
    pg_qrnote = doc.add_paragraph(); pg_qrnote.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = pg_qrnote.add_run('扫码关注微信公众号')
    r.font.size = Pt(8); r.font.color.rgb = GRAY; r.font.name = '微软雅黑'
    if r._element.rPr is None: r._element.rPr = r._element.get_or_add_rPr()
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_page_break()

    # ========== 一、执行摘要 ==========
    h1(doc, '一', '执行摘要')
    note(doc, '资本金IRR=使资本金净现金流序列NPV=0的折现率（基准=WACC=5.5%%，资本金IRR=%.2f%%）' % base_irr)
    mets = [
        ('全投资 IRR', '%.2f%%' % m['irr_full'], '全投资折现率 WACC=5.5%%'),
        ('资本金 IRR', '%.2f%%' % m['irr_equity'], '资本金折现率 equity cost=8%%'),
        ('全投资 NPV', '%.0f 万元' % m['npv_full'], 'WACC=5.5%% 折现'),
        ('资本金 NPV', '%.0f 万元' % m['npv_equity'], '资本金成本=8%% 折现'),
        ('ROE（年均）', '%.1f%%' % m['roe'], '年均净利润/资本金'),
        ('ROA（年均）', '%.2f%%' % m['roa'], '年均EBIT/总资产'),
        ('ROIC（年均）', '%.2f%%' % m['roic'], 'NOPAT/(债务+权益)'),
        ('ROI（累计）', '%.1f%%' % m['roi'], '累计净收入/总投资'),
        ('年发电量', '%.0f 万kWh' % p['annual_power'], '利用小时=%dh，限电率=%.0f%%' % (p['util_hours'], p['curtailment']*100)),
        ('年利用小时', '%d h' % p['util_hours'], '限电%.0f%%后有效小时' % (p['curtailment']*100)),
        ('公司投资', '%.2f 亿元' % (p['total_inv']/10000), 'EPC单价=%.1f元/W' % p['epc_price']),
        ('资本金', '%.0f 万元' % p['capital'], '资本金比例=20%%（融资杠杆=0.8）'),
        ('贷款利率', '3.5%%', '含税'),
        ('融资期限', '%d 年' % p['loan_years'], '等额本金'),
        ('残值', '%.0f 万元' % p['residual_value'], '总投资×3%%'),
        ('年折旧', '%.0f 万元/年' % p['annual_depreciation'], '折旧基数=总投资×97%%'),
    ]
    make_table(doc, ['指标', '数值', '说明'], [(k,v,d) for k,v,d in mets], [4.0, 3.5, 7.0], sz=9)

    # ========== 二~五：各类表格 ==========
    # 二、全投资CF
    h1(doc, '二', '全投资现金流量表（全生命周期）')
    note(doc, '全投资CF = 发电收入(含税) + 进项税抵扣 - 运营成本 - 销项税 - 城建税及附加 - 调整所得税')
    hdrs2 = ['年', '电价', '发电收入', '进项税抵扣', '运营成本', '销项税', '城建税\n及附加', '调整所得税', '本金', '利息', '净现金流']
    rows2 = [(str(d['year']), '%.3f'%d['price'], '%.1f'%d['revenue_vat'], '%.1f'%d['deduct_vat'],
              '%.1f'%d['op_cost'], '%.1f'%d['vat_output'], '%.1f'%d['city_edu_tax'],
              '%.1f'%d['adj_tax'], '%.1f'%d['principal'], '%.1f'%d['interest'], '%.1f'%d['cf_full'])
             for d in D['yearly']]
    make_table(doc, hdrs2, rows2, [0.8, 1.2, 2.2, 2.2, 2.2, 2.0, 2.2, 2.2, 2.0, 2.0, 2.4], sz=7.5)

    # 三、资本金CF
    h1(doc, '三', '资本金现金流量表（全生命周期）')
    note(doc, '资本金CF = 发电收入(含税) + 进项税抵扣 - 运营成本 - 销项税 - 城建税及附加 - 所得税 - 本金 - 利息')
    hdrs3 = ['年', '发电收入', '进项税抵扣', '运营成本', '销项税', '城建税\n及附加', '所得税', '本金', '利息', '净现金流']
    rows3 = [(str(d['year']), '%.1f'%d['revenue_vat'], '%.1f'%d['deduct_vat'], '%.1f'%d['op_cost'],
              '%.1f'%d['vat_output'], '%.1f'%d['city_edu_tax'], '%.1f'%d['income_tax'],
              '%.1f'%d['principal'], '%.1f'%d['interest'], '%.1f'%d['cf_equity'])
             for d in D['yearly']]
    make_table(doc, hdrs3, rows3, [0.8, 2.5, 2.5, 2.5, 2.2, 2.5, 2.2, 2.2, 2.2, 2.5], sz=7.5)

    # 四、利润表
    h1(doc, '四', '资本金利润表（全生命周期）')
    note(doc, '净利润 = 发电收入(不含税) - 运维成本 - 保险备件 - 折旧 - 利息（PBT>0时x25%%所得税）')
    hdrs4 = ['年', '发电收入\n(不含税)', '运维成本', '保险备件', '折旧', '利息', 'PBT', '所得税', '净利润']
    rows4 = [(str(d['year']), '%.1f'%d['revenue_net'], '%.1f'%d['om_cost'], '%.1f'%(d['insurance']+d['spare']),
              '%.1f'%d['depreciation'], '%.1f'%d['interest'], '%.1f'%d['pbt'],
              '%.1f'%d['income_tax'], '%.1f'%d['net_profit'])
             for d in D['yearly']]
    make_table(doc, hdrs4, rows4, [0.8, 2.5, 2.2, 2.0, 2.0, 2.0, 2.2, 2.0, 2.2], sz=7.5)

    # 五、FCFE
    h1(doc, '五', '资本金自由现金流量表（FCFE）')
    note(doc, 'FCFE = 净利润 + 折旧 + 利息抵税(x75%%) - 本金支出')
    hdrs5 = ['年', '净利润', '折旧', '利息(抵税后)', '本金支出', 'FCFE']
    rows5 = [(str(d['year']), '%.1f'%d['net_profit'], '%.1f'%d['depreciation'],
              '%.1f'%(d['interest']*(1-0.25)), '%.1f'%d['principal'], '%.1f'%d['fcfe'])
             for d in D['yearly']]
    make_table(doc, hdrs5, rows5, [0.8, 3.0, 2.8, 3.0, 2.8, 3.0], sz=7.5)

    # 六、DSCR
    h1(doc, '六', '偿债覆盖率（DSCR）')
    note(doc, 'DSCR = (净利润 + 折旧 + 利息) / (本金 + 利息)  |  >1.2 健康  |  1.0~1.2 预警  |  <1.0 危险')
    rows6 = [(str(d['year']), '%.3f'%d['dscr'] if d['dscr'] is not None else '已还清', dscr_status(d['dscr']))
             for d in D['yearly']]
    make_table(doc, ['年度', 'DSCR', '偿债状态'], rows6, [2.0, 3.5, 3.5], sz=9)

    # 七、敏感性分析
    h1(doc, '七', '敏感性分析')
    note(doc, '基准资本金IRR=%.2f%%，以下为各因素变动对资本金IRR的影响' % base_irr)
    for key, title in [
        ('cfd_price', '电价对资本金IRR的影响'),
        ('util_hours', '利用小时数对资本金IRR的影响'),
        ('total_inv', '总投资对资本金IRR的影响'),
        ('loan_rate', '贷款利率对资本金IRR的影响'),
        ('om_cost', '运维成本对资本金IRR的影响'),
    ]:
        h2(doc, title)
        if key == 'loan_rate':
            rows_sens = [('%.1fpp(%+dbp)'%(d/100,int(d)), '%.2f%%'%irr, '%+.2f%%'%(irr-base_irr))
                         for d, irr in D['sens'][key]]
        else:
            rows_sens = [('%+.0f%%'%d, '%.2f%%'%irr, '%+.2f%%'%(irr-base_irr))
                         for d, irr in D['sens'][key]]
        make_table(doc, ['变动', '资本金IRR', 'IRR变化'], rows_sens, [4.5, 3.0, 3.0], sz=9)

    # 八、关键公式
    h1(doc, '八', '关键计算公式')
    formulas = [
        ('总投资', 'EPC单价x装机容量 = %.1f元/W x %dMW = %.0f万元' % (p['epc_price'], p['capacity'], p['total_inv'])),
        ('资本金', '总投资x20%% = %.0f万元' % p['capital']),
        ('贷款', '总投资x80%% = %.0f万元' % p['loan']),
        ('年发电量', '60MW x %dh x (1-%.0f%%) x 1000 = %.0f万kWh' % (p['util_hours'], p['curtailment']*100, p['annual_power'])),
        ('折旧基数', '总投资x97%% = %.0f万元' % (p['total_inv']*0.97)),
        ('年折旧', '折旧基数/20年 = %.0f万元/年' % p['annual_depreciation']),
        ('残值', '总投资x3%% = %.0f万元' % p['residual_value']),
        ('进项税抵扣', '(年本金+年利息)/1.13x13%%，仅在融资期限内'),
        ('调整所得税', '(收入不含税-运营成本-折旧)x25%%（全投资CF用，不加回利息）'),
        ('DSCR', '(净利润+折旧+利息)/(本金+利息)'),
        ('全投资IRR', 'IRR([-总投资] + [全投资CF序列]) = %.2f%%' % m['irr_full']),
        ('资本金IRR', 'IRR([-资本金] + [资本金CF序列]) = %.2f%%' % m['irr_equity']),
    ]
    make_table(doc, ['项目', '公式/结果'], formulas, [4.0, 10.5])

    # 九、分析图表
    h1(doc, '九', '分析图表')
    CP = [
        ('chart1.png', '图1：全投资净现金流趋势（万元）'),
        ('chart2.png', '图2：资本金净现金流趋势（万元）'),
        ('chart3.png', '图3：累计现金流对比（万元）'),
        ('chart4.png', '图4：收入成本结构（全生命周期）'),
        ('chart5.png', '图5：DSCR偿债覆盖率'),
        ('chart6.png', '图6：资本金IRR敏感性（电价%）'),
        ('chart7.png', '图7：多变量敏感性分析'),
    ]
    for fname, title in CP:
        path = os.path.join(charts_dir, fname)
        if os.path.exists(path):
            h2(doc, title)
            doc.add_picture(path, width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(output_path)
    size_kb = os.path.getsize(output_path) // 1024
    print(f'Word saved: {output_path} ({size_kb}KB)')


if __name__ == '__main__':
    if len(sys.argv) < 6:
        print("用法: python gen_wind_report.py <json_path> <charts_dir> <logo_path> <qr_path> <output_docx>")
        sys.exit(1)
    generate_report(sys.argv[1], sys.argv[2], sys.argv[3], sys.argv[4], sys.argv[5])
